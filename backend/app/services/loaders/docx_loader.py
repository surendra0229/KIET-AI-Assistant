from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import LoadedChunk, clean_text


def load_docx(path: Path) -> list[LoadedChunk]:
    try:
        doc = Document(str(path))
    except PackageNotFoundError as e:
        raise ValueError(f"Corrupted or invalid DOCX: {path.name}") from e

    chunks: list[LoadedChunk] = []
    para_no = 0
    for para in doc.paragraphs:
        text = clean_text(para.text or "")
        if not text:
            continue
        para_no += 1
        chunks.append(
            LoadedChunk(
                text=text,
                metadata={
                    "document_name": path.name,
                    "paragraph_number": para_no,
                    "source_type": "docx",
                },
            )
        )

    # Include tables too — many college docs (schedules, fee slabs) live there.
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows, start=1):
            cells = [clean_text(c.text) for c in row.cells]
            row_text = " | ".join(c for c in cells if c)
            if not row_text:
                continue
            para_no += 1
            chunks.append(
                LoadedChunk(
                    text=row_text,
                    metadata={
                        "document_name": path.name,
                        "paragraph_number": para_no,
                        "table_index": t_idx,
                        "row_index": r_idx,
                        "source_type": "docx-table",
                    },
                )
            )
    return chunks
